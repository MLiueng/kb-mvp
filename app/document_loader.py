"""文档解析模块：根据文件扩展名自动选择解析器，提取带位置信息的结构化内容。

输出 ParsedDocument 包含全文文本与按段落/页码切分的 blocks，每个 block 携带
char_start / char_end / page_number / para_index 等位置信息，为答案溯源提供精确定位基础。
"""
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import List, Optional, Tuple, Iterator


@dataclass
class Block:
    """结构化内容块（解析阶段产物）。"""
    block_index: int
    block_type: str            # paragraph / heading
    page_number: Optional[int]  # PDF 页码，其他类型为 None
    para_index: Optional[int]   # 段落序号
    char_start: int            # 全文起始字符偏移
    char_end: int              # 全文结束字符偏移
    text: str


@dataclass
class ParsedDocument:
    """解析后的文档结构。"""
    full_text: str
    file_type: str             # pdf / docx / txt / md
    total_pages: Optional[int]
    blocks: List[Block] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "full_text": self.full_text,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "blocks": [asdict(b) for b in self.blocks],
        }

    @classmethod
    def from_dict(cls, d: dict) -> "ParsedDocument":
        return cls(
            full_text=d["full_text"],
            file_type=d["file_type"],
            total_pages=d.get("total_pages"),
            blocks=[Block(**b) for b in d.get("blocks", [])],
        )


# 解析器注册表：扩展名 -> 解析函数
_SEPARATOR = "\n\n"


def load_document(file_path: str) -> ParsedDocument:
    """根据扩展名选择解析器，返回 ParsedDocument。"""
    p = Path(file_path)
    if not p.exists():
        raise ValueError(f"文件不存在: {file_path}")
    ext = p.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(p)
    if ext in (".docx", ".doc"):
        return _parse_docx(p)
    if ext == ".txt":
        return _parse_text(p, "txt")
    if ext in (".md", ".markdown"):
        return _parse_text(p, "md")
    raise ValueError(f"不支持的文件格式: {ext}（支持 pdf/docx/txt/md）")


def _assemble(block_iter: Iterator[Tuple[Optional[int], Optional[int], str, str]],
              file_type: str, total_pages: Optional[int]) -> ParsedDocument:
    """将 (page, para_idx, text, block_type) 流组装为 ParsedDocument。

    全文由各 block 文本以 \\n\\n 连接而成，char_start/char_end 与 full_text 共享同一坐标系。
    """
    blocks: List[Block] = []
    parts: List[str] = []
    char_offset = 0
    block_index = 0
    for page_number, para_index, text, block_type in block_iter:
        if not text:
            continue
        text = text.strip()
        if not text:
            continue
        start = char_offset
        end = start + len(text)
        blocks.append(Block(
            block_index=block_index,
            block_type=block_type,
            page_number=page_number,
            para_index=para_index,
            char_start=start,
            char_end=end,
            text=text,
        ))
        parts.append(text)
        parts.append(_SEPARATOR)
        char_offset = end + len(_SEPARATOR)
        block_index += 1
    full_text = "".join(parts)
    if not full_text.strip():
        raise ValueError("文档内容为空，无法入库")
    return ParsedDocument(
        full_text=full_text,
        file_type=file_type,
        total_pages=total_pages,
        blocks=blocks,
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    """PyMuPDF 逐页提取文本，按段落拆分，记录页码与字符偏移。"""
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    total_pages = doc.page_count

    def iter_blocks():
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if not text:
                continue
            # 按空行切分段落
            paragraphs = [p for p in text.split("\n\n") if p.strip()]
            for para_idx, para in enumerate(paragraphs):
                yield page_num, para_idx, para, _guess_block_type(para)
        doc.close()

    return _assemble(iter_blocks(), "pdf", total_pages)


def _parse_docx(path: Path) -> ParsedDocument:
    """python-docx 提取段落，记录段落序号与字符偏移。"""
    from docx import Document

    document = Document(str(path))

    def iter_blocks():
        for para_idx, para in enumerate(document.paragraphs):
            text = para.text
            if not text or not text.strip():
                continue
            yield None, para_idx, text, _guess_block_type(text)

    # python-docx 难以直接获取页码，total_pages 置 None
    return _assemble(iter_blocks(), "docx", None)


def _parse_text(path: Path, file_type: str) -> ParsedDocument:
    """UTF-8 直接读取纯文本/Markdown，按空行切分段落。"""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    def iter_blocks():
        # 优先按空行切分；保留 Markdown 标题作为独立 block
        paragraphs = [p for p in text.split("\n\n") if p.strip()]
        for para_idx, para in enumerate(paragraphs):
            yield None, para_idx, para, _guess_block_type(para)

    return _assemble(iter_blocks(), file_type, None)


def _guess_block_type(text: str) -> str:
    """启发式判断 block 类型：Markdown 标题识别为 heading。"""
    stripped = text.lstrip()
    if stripped.startswith("#"):
        return "heading"
    return "paragraph"
